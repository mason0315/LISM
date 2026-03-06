import base64
import io
from docx import Document
import json

def extract_text_from_base64_docx(base64_data):
    """
    从base64编码的docx数据中提取文本内容
    """
    try:
        # 提取base64数据部分（去掉data:application/...;base64, 前缀）
        if ',' in base64_data:
            base64_data = base64_data.split(',')[1]
        
        # 解码base64数据
        decoded_data = base64.b64decode(base64_data)
        
        # 创建文件流
        file_stream = io.BytesIO(decoded_data)
        
        # 读取Word文档
        doc = Document(file_stream)
        
        # 提取文本
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # 对于表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        return '\n'.join(text_content)
        
    except Exception as e:
        return f"处理文档时出错: {str(e)}"

if __name__ == "__main__":
    # 示例用法
    print("此脚本用于从base64编码的docx文档中提取文本内容")
    print("使用方法: import process_docx; text = process_docx.extract_text_from_base64_docx(base64_string)")